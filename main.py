# main.py
from fastapi import FastAPI, Request, Form
from fastapi.responses import JSONResponse, HTMLResponse, FileResponse
from pydantic import BaseModel
import feedparser
import os
import json
from datetime import datetime
import hashlib
import bleach
from pathlib import Path

DATA_DIR = Path.cwd() / "data"
FEEDS_FILE = DATA_DIR / "feeds.json"

app = FastAPI(title="ReadNest RSS Ingest (dev)")

# Ensure data dir + file exist
def ensure_data_file():
    DATA_DIR.mkdir(exist_ok=True)
    if not FEEDS_FILE.exists():
        FEEDS_FILE.write_text(json.dumps({"feeds": []}, indent=2), encoding="utf-8")

def read_db():
    ensure_data_file()
    return json.loads(FEEDS_FILE.read_text(encoding="utf-8"))

def write_db(obj):
    tmp = FEEDS_FILE.with_suffix(".tmp")
    tmp.write_text(json.dumps(obj, indent=2), encoding="utf-8")
    tmp.replace(FEEDS_FILE)

def hash_url(url: str) -> str:
    return hashlib.sha256(url.encode("utf-8")).hexdigest()[:12]

def sanitize_html(html: str) -> str:
    # Bleach to allow a conservative set of tags; keep it safe.
    allowed_tags = ["p","br","strong","em","ul","ol","li","a","blockquote","code","pre"]
    allowed_attrs = {"a":["href","title","rel"], "img":["src","alt"]}
    return bleach.clean(html or "", tags=allowed_tags, attributes=allowed_attrs, strip=True)

class IngestPayload(BaseModel):
    url: str

@app.get("/", response_class=HTMLResponse)
async def index():
    # serve the simple prompt UI
    html = (Path(__file__).parent / "static" / "index.html").read_text(encoding="utf-8")
    return HTMLResponse(content=html)

@app.post("/api/ingest-rss")
async def ingest_rss(payload: IngestPayload):
    url = payload.url.strip()
    if not url:
        return JSONResponse({"error": "missing url"}, status_code=400)

    try:
        parsed = feedparser.parse(url)
    except Exception as e:
        return JSONResponse({"error": "failed to fetch or parse feed", "details": str(e)}, status_code=500)

    if parsed.bozo and not parsed.entries:
        # bozo flag indicates parse issues; still try to be helpful
        return JSONResponse({"error": "invalid RSS/Atom or unreachable", "details": getattr(parsed, 'bozo_exception', str(parsed))}, status_code=400)

    items = []
    for it in parsed.entries:
        content = ""
        # prefer content:encoded or content
        if "content" in it and isinstance(it.content, list) and len(it.content) > 0:
            content = it.content[0].value
        elif "summary" in it:
            content = it.summary
        elif "description" in it:
            content = it.description

        items.append({
            "guid": it.get("id") or it.get("guid") or None,
            "title": bleach.clean(it.get("title","Untitled"), strip=True),
            "link": it.get("link"),
            "content": sanitize_html(content),
            "published": it.get("published") or it.get("updated") or None,
            "isoDate": it.get("published_parsed") and datetime(*it.published_parsed[:6]).isoformat() if it.get("published_parsed") else None,
            "author": it.get("author") or it.get("creator") or None
        })

    stored_feed = {
        "id": hash_url(url),
        "url": url,
        "title": parsed.feed.get("title") if parsed.feed else None,
        "fetchedAt": datetime.utcnow().isoformat() + "Z",
        "items": items
    }

    db = read_db()
    # replace by id (url hash) if exists
    existing_idx = next((i for i,f in enumerate(db["feeds"]) if f["id"] == stored_feed["id"]), None)
    if existing_idx is not None:
        db["feeds"][existing_idx] = stored_feed
    else:
        db["feeds"].append(stored_feed)

    write_db(db)
    return {"ok": True, "storedCount": len(items), "feedTitle": stored_feed["title"]}

@app.get("/api/feeds")
async def list_feeds():
    return read_db()

@app.get("/api/feeds/{feed_id}")
async def get_feed(feed_id: str):
    db = read_db()
    feed = next((f for f in db["feeds"] if f["id"] == feed_id), None)
    if not feed:
        return JSONResponse({"error":"not found"}, status_code=404)
    return feed
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Query, Depends, Header
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from scholar_agent import scholar_agent
from hybrid_service import hybrid_service
import json
import os
from datetime import datetime
import uuid
import feedparser
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
import io
import httpx
import traceback, json
from supabase_config import supabase

app = FastAPI(title="ReadNest Backend", version="1.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class JournalEntry(BaseModel):
    id: str
    title: str
    content: str
    created_at: str
    updated_at: str
    word_count: int
    keywords: Optional[dict] = {}
    user_id: Optional[str] = None

class JournalCreate(BaseModel):
    title: str
    content: str = ""

class JournalUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None

class RegisterRequest(BaseModel):
    email: str
    password: str

class LoginRequest(BaseModel):
    email: str
    password: str

def get_current_user(authorization: Optional[str] = Header(None)):
    """Validate Supabase JWT from Authorization: Bearer <token> and return user dict."""
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")
    token = authorization.split(" ", 1)[1].strip()
    try:
        # Ensure subsequent PostgREST calls run under this user's RLS context
        try:
            supabase.postgrest.auth(token)
        except Exception:
            pass
        user_resp = supabase.auth.get_user(token)
        user = user_resp.user
        if not user:
            raise HTTPException(status_code=401, detail="Invalid token")
        return {"id": user.id, "email": user.email}
    except Exception:
        raise HTTPException(status_code=401, detail="Invalid token")

# Using hybrid service (Supabase + JSON fallback)

def extract_keywords(text: str, top_n: int = 30) -> dict:
    """Extract meaningful keywords from text, filtering out common stop words"""
    if not text:
        return {}
    
    import re
    
    # Common stop words to filter out
    stop_words = {
        'a', 'an', 'and', 'are', 'as', 'at', 'be', 'been', 'by', 'for', 'from',
        'has', 'he', 'in', 'is', 'it', 'its', 'of', 'on', 'that', 'the', 'to',
        'was', 'were', 'will', 'with', 'the', 'this', 'but', 'they', 'have',
        'had', 'what', 'said', 'each', 'which', 'their', 'time', 'will', 'about',
        'if', 'up', 'out', 'many', 'then', 'them', 'these', 'so', 'some', 'her',
        'would', 'make', 'like', 'into', 'him', 'has', 'two', 'more', 'write',
        'go', 'see', 'number', 'no', 'way', 'could', 'people', 'my', 'than',
        'first', 'been', 'call', 'who', 'oil', 'sit', 'now', 'find', 'down',
        'day', 'did', 'get', 'come', 'made', 'may', 'part', 'new', 'save',
        'entry', 'entries', 'note', 'notes', 'journal', 'journals', 'read',
        'reading', 'text', 'content', 'title', 'keyword', 'keywords'
    }
    
    # Tokenize and clean text
    tokens = re.sub(r'[^\w\s]', ' ', text.lower()).split()
    
    # Filter: keep only words that are:
    # - Longer than 3 characters (more meaningful)
    # - Not in stop words list
    # - Not pure numbers
    tokens = [
        t for t in tokens 
        if len(t) > 3 
        and t not in stop_words 
        and not t.isdigit()
    ]
    
    if not tokens:
        return {}
    
    # Count frequency
    freq = {}
    for token in tokens:
        freq[token] = freq.get(token, 0) + 1
    
    # Return top N keywords sorted by frequency
    sorted_keywords = sorted(freq.items(), key=lambda x: x[1], reverse=True)[:top_n]
    return dict(sorted_keywords)

@app.get("/")
def read_root():
    return {"message": "ReadNest Backend API", "version": "1.0.0"}

# Journal API endpoints
@app.get("/api/journals", response_model=List[JournalEntry])
def get_all_journals(current_user: dict = Depends(get_current_user)):
    """Get all journal entries"""
    return hybrid_service.get_all_journals(user_id=current_user["id"])

@app.get("/api/journals/{journal_id}", response_model=JournalEntry)
def get_journal(journal_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific journal entry"""
    journal = hybrid_service.get_journal(journal_id)
    if not journal:
        raise HTTPException(status_code=404, detail="Journal not found")
    # RLS prevents cross-user reads from DB; add explicit check for JSON fallback
    if getattr(journal, "user_id", None) and journal.user_id != current_user["id"]:
        raise HTTPException(status_code=404, detail="Journal not found")
    return journal

@app.post("/api/journals", response_model=JournalEntry)
def create_journal(journal_data: JournalCreate, current_user: dict = Depends(get_current_user)):
    """Create a new journal entry"""
    # Calculate word count and extract keywords
    word_count = len(journal_data.content.split()) if journal_data.content else 0
    keywords = extract_keywords(journal_data.content)
    
    journal_dict = {
        'title': journal_data.title,
        'content': journal_data.content,
        'word_count': word_count,
        'keywords': keywords,
        'user_id': current_user["id"]
    }
    
    new_journal = hybrid_service.create_journal(journal_dict)
    if not new_journal:
        raise HTTPException(status_code=500, detail="Failed to create journal")
    
    return new_journal

@app.put("/api/journals/{journal_id}", response_model=JournalEntry)
def update_journal(journal_id: str, journal_data: JournalUpdate, current_user: dict = Depends(get_current_user)):
    """Update an existing journal entry"""
    update_dict = {}
    
    # Update fields if provided
    if journal_data.title is not None:
        update_dict['title'] = journal_data.title
    if journal_data.content is not None:
        update_dict['content'] = journal_data.content
        update_dict['word_count'] = len(journal_data.content.split())
        update_dict['keywords'] = extract_keywords(journal_data.content)
    
    updated_journal = hybrid_service.update_journal(journal_id, update_dict)
    if not updated_journal:
        raise HTTPException(status_code=404, detail="Journal not found")
    if getattr(updated_journal, "user_id", None) and updated_journal.user_id != current_user["id"]:
        raise HTTPException(status_code=403, detail="Forbidden")
    
    return updated_journal

@app.delete("/api/journals/{journal_id}")
def delete_journal(journal_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a journal entry"""
    success = hybrid_service.delete_journal(journal_id)
    if not success:
        raise HTTPException(status_code=404, detail="Journal not found")
    
    return {"message": "Journal deleted successfully"}

@app.get("/api/journals/search/{query}")
def search_journals(query: str, current_user: dict = Depends(get_current_user)):
    """Search journals by title or content"""
    return hybrid_service.search_journals(query, user_id=current_user["id"])

# RSS Feed System

class FeedSubscription(BaseModel):
    id: str
    url: str
    title: str
    description: str
    last_updated: str
    is_active: bool = True
    user_id: Optional[str] = None

class Article(BaseModel):
    id: str
    title: str
    source: str
    snippet: str
    date: str
    type: str  # 'rss' or 'pdf'
    url: Optional[str] = None
    feed_id: Optional[str] = None
    content: Optional[str] = None
    author: Optional[str] = None
    tags: Optional[List[str]] = []
    user_id: Optional[str] = None

class FeedCreate(BaseModel):
    url: str
    name: str

class Document(BaseModel):
    id: str
    name: str
    type: str  # 'pdf' or 'doc'
    size: int
    upload_date: str
    content: Optional[str] = None
    status: str = "ready"  # 'uploading', 'processing', 'ready', 'error'
    user_id: Optional[str] = None

# Storage files
FEEDS_FILE = "feed_subscriptions.json"
ARTICLES_FILE = "articles.json"
DOCUMENTS_FILE = "documents.json"

def load_feed_subscriptions() -> List[FeedSubscription]:
    """Load feed subscriptions from local JSON file"""
    if os.path.exists(FEEDS_FILE):
        try:
            with open(FEEDS_FILE, 'r') as f:
                data = json.load(f)
                return [FeedSubscription(**feed) for feed in data]
        except Exception:
            return []
    return []

def save_feed_subscriptions(subscriptions: List[FeedSubscription]):
    """Save feed subscriptions to local JSON file"""
    with open(FEEDS_FILE, 'w') as f:
        json.dump([sub.dict() for sub in subscriptions], f, indent=2)

def load_articles() -> List[Article]:
    """Load articles from local JSON file"""
    if os.path.exists(ARTICLES_FILE):
        try:
            with open(ARTICLES_FILE, 'r') as f:
                data = json.load(f)
                return [Article(**article) for article in data]
        except Exception:
            return []
    return []

def save_articles(articles: List[Article]):
    """Save articles to local JSON file"""
    with open(ARTICLES_FILE, 'w') as f:
        json.dump([article.dict() for article in articles], f, indent=2)

def load_documents() -> List[Document]:
    """Load documents from local JSON file"""
    if os.path.exists(DOCUMENTS_FILE):
        try:
            with open(DOCUMENTS_FILE, 'r') as f:
                data = json.load(f)
                return [Document(**doc) for doc in data]
        except Exception:
            return []
    return []

def save_documents(documents: List[Document]):
    """Save documents to local JSON file"""
    with open(DOCUMENTS_FILE, 'w') as f:
        json.dump([doc.dict() for doc in documents], f, indent=2)

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def process_document(file: UploadFile) -> Document:
    """Process uploaded document and extract text"""
    # Read file content
    file_content = file.file.read()
    
    # Determine file type
    file_type = "pdf" if file.content_type == "application/pdf" else "doc"
    
    # Extract text based on file type
    try:
        if file_type == "pdf":
            content = extract_text_from_pdf(file_content)
        else:  # docx
            content = extract_text_from_docx(file_content)
    except Exception as e:
        raise Exception(f"Failed to process document: {str(e)}")
    
    # Create document object
    document = Document(
        id=f"doc_{int(datetime.now().timestamp())}_{uuid.uuid4().hex[:8]}",
        name=file.filename,
        type=file_type,
        size=len(file_content),
        upload_date=datetime.now().isoformat(),
        content=content,
        status="ready"
    )
    
    return document

def parse_rss_feed(feed_url: str, custom_name: str = None) -> tuple[FeedSubscription, List[Article]]:
    """Parse RSS feed and return subscription info and articles"""
    try:
        # Parse the RSS feed
        feed = feedparser.parse(feed_url)
        
        if feed.bozo:
            raise Exception(f"Invalid RSS feed: {feed.bozo_exception}")
        
        if not feed.entries:
            raise Exception("No entries found in RSS feed")
        
        # Create feed subscription
        subscription = FeedSubscription(
            id=f"feed_{int(datetime.now().timestamp())}_{uuid.uuid4().hex[:8]}",
            url=feed_url,
            title=custom_name or feed.feed.get('title', 'Untitled Feed'),
            description=feed.feed.get('description', 'No description'),
            last_updated=datetime.now().isoformat(),
            is_active=True
        )
        
        # Parse articles
        articles = []
        for entry in feed.entries[:50]:  # Limit to 50 most recent articles
            # Extract content
            content = ""
            if hasattr(entry, 'content') and entry.content:
                content = entry.content[0].value if entry.content else ""
            elif hasattr(entry, 'summary'):
                content = entry.summary
            
            # Clean HTML content
            if content:
                soup = BeautifulSoup(content, 'html.parser')
                content = soup.get_text()[:1000]  # Limit to 1000 chars
            
            # Extract snippet
            snippet = content[:200] + "..." if len(content) > 200 else content
            if not snippet and hasattr(entry, 'summary'):
                snippet = entry.summary[:200] + "..." if len(entry.summary) > 200 else entry.summary
            
            # Parse date
            article_date = datetime.now().isoformat().split('T')[0]  # Default to today
            if hasattr(entry, 'published_parsed') and entry.published_parsed:
                try:
                    article_date = datetime(*entry.published_parsed[:6]).isoformat().split('T')[0]
                except:
                    pass
            
            # Extract tags properly
            tags = []
            if hasattr(entry, 'tags') and entry.tags:
                for tag in entry.tags[:5]:  # Limit to 5 tags
                    if hasattr(tag, 'term') and tag.term:
                        tags.append(tag.term)
                    elif isinstance(tag, str):
                        tags.append(tag)
            
            article = Article(
                id=f"article_{int(datetime.now().timestamp())}_{uuid.uuid4().hex[:8]}",
                title=entry.get('title', 'Untitled Article'),
                source=subscription.title,
                snippet=snippet,
                date=article_date,
                type="rss",
                url=entry.get('link', ''),
                feed_id=subscription.id,
                content=content,
                author=entry.get('author', ''),
                tags=tags
            )
            articles.append(article)
        
        return subscription, articles
        
    except Exception as e:
        raise Exception(f"Failed to parse RSS feed: {str(e)}")

def refresh_all_feeds():
    """Refresh all active feed subscriptions"""
    subscriptions = load_feed_subscriptions()
    all_articles = load_articles()
    
    for subscription in subscriptions:
        if not subscription.is_active:
            continue
            
        try:
            _, new_articles = parse_rss_feed(subscription.url)
            
            # Update subscription
            subscription.last_updated = datetime.now().isoformat()
            
            # Add new articles (avoid duplicates)
            existing_urls = {article.url for article in all_articles if article.url}
            for article in new_articles:
                if article.url not in existing_urls:
                    all_articles.append(article)
            
        except Exception as e:
            print(f"Failed to refresh feed {subscription.title}: {e}")
    
    # Save updated data
    save_feed_subscriptions(subscriptions)
    save_articles(all_articles)
    
    return all_articles

# Feed API endpoints
@app.get("/api/feeds", response_model=List[Article])
def get_all_articles(current_user: dict = Depends(get_current_user)):
    """Get all articles for the current user from active feeds"""
    articles = hybrid_service.get_all_articles(user_id=current_user["id"])
    articles.sort(key=lambda x: x.date, reverse=True)
    return articles

@app.get("/api/feeds/subscriptions", response_model=List[FeedSubscription])
def get_feed_subscriptions(current_user: dict = Depends(get_current_user)):
    """Get all feed subscriptions for the current user"""
    return hybrid_service.get_all_feed_subscriptions(user_id=current_user["id"])

@app.post("/api/feeds", response_model=dict)
def add_feed_subscription(feed_data: FeedCreate, current_user: dict = Depends(get_current_user)):
    """Add a new RSS feed subscription"""
    try:
        # Parse the RSS feed
        subscription, articles = parse_rss_feed(feed_data.url, feed_data.name)
        
        # Associate with user and persist via database-first hybrid service
        subscription.user_id = current_user["id"]
        for a in articles:
            a.user_id = current_user["id"]
        created_sub = hybrid_service.create_feed_subscription(subscription.dict())
        for a in articles:
            hybrid_service.create_article(a.dict())
        
        return {
            "message": "Feed added successfully",
            "subscription": created_sub or subscription,
            "articles_count": len(articles)
        }
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.delete("/api/feeds/subscriptions/{subscription_id}")
def delete_feed_subscription(subscription_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a feed subscription and its articles"""
    # DB will enforce ownership with RLS; we return generic 200 on success
    hybrid_service.delete_feed_subscription(subscription_id)
    return {"message": "Feed subscription deleted successfully"}

@app.post("/api/feeds/subscriptions/{subscription_id}/toggle")
def toggle_feed_subscription(subscription_id: str, toggle_data: dict, current_user: dict = Depends(get_current_user)):
    """Toggle feed subscription active status"""
    # Fetch, then update via database
    subs = hybrid_service.get_all_feed_subscriptions(user_id=current_user["id"])
    subscription = next((s for s in subs if s.id == subscription_id), None)
    if not subscription:
        raise HTTPException(status_code=404, detail="Feed subscription not found")
    subscription.is_active = toggle_data.get('is_active', not subscription.is_active)
    hybrid_service.create_feed_subscription(subscription.dict())  # upsert behavior via hybrid path
    return {
        "message": f"Feed subscription {'activated' if subscription.is_active else 'deactivated'} successfully",
        "is_active": subscription.is_active
    }

@app.post("/api/feeds/refresh")
def refresh_feeds(current_user: dict = Depends(get_current_user)):
    """Manually refresh all feeds for the current user"""
    try:
        # Re-parse each active feed for this user and store new articles
        subscriptions = hybrid_service.get_all_feed_subscriptions(user_id=current_user["id"])
        total = 0
        for sub in [s for s in subscriptions if s.is_active]:
            _, new_articles = parse_rss_feed(sub.url, sub.title)
            for a in new_articles:
                a.user_id = current_user["id"]
                a.feed_id = sub.id
                hybrid_service.create_article(a.dict())
                total += 1
        return {"message": "Feeds refreshed successfully", "total_articles": total}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to refresh feeds: {str(e)}")

@app.get("/api/feeds/search/{query}")
def search_articles(query: str, current_user: dict = Depends(get_current_user)):
    """Search articles by title, content, or tags for current user"""
    articles = hybrid_service.get_all_articles(user_id=current_user["id"])
    query_lower = query.lower()
    
    matching_articles = []
    for article in articles:
        if (query_lower in article.title.lower() or 
            query_lower in article.snippet.lower() or
            query_lower in (article.content or "").lower() or
            any(query_lower in tag.lower() for tag in (article.tags or []))):
            matching_articles.append(article)
    
    return matching_articles

# Document API endpoints
@app.get("/api/documents", response_model=List[Document])
def get_all_documents(current_user: dict = Depends(get_current_user)):
    """Get all uploaded documents for the current user"""
    documents = hybrid_service.get_all_documents(user_id=current_user["id"])
    documents.sort(key=lambda x: x.upload_date, reverse=True)
    return documents

@app.get("/api/documents/{document_id}", response_model=Document)
def get_document(document_id: str, current_user: dict = Depends(get_current_user)):
    """Get a specific document"""
    documents = hybrid_service.get_all_documents(user_id=current_user["id"])
    document = next((d for d in documents if d.id == document_id), None)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    return document

@app.post("/api/documents/upload", response_model=Document)
def upload_document(file: UploadFile = File(...), name: str = Form(...), current_user: dict = Depends(get_current_user)):
    """Upload and process a document"""
    # Validate file type
    if file.content_type not in ["application/pdf", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    
    # Validate file size (10MB limit)
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Reset to beginning
    
    if file_size > 10 * 1024 * 1024:  # 10MB
        raise HTTPException(status_code=400, detail="File size must be less than 10MB")
    
    try:
        # Process the document
        document = process_document(file)
        
        # Persist via database-first hybrid service with user scoping
        created = hybrid_service.create_document({
            'name': document.name,
            'type': document.type,
            'size': document.size,
            'content': document.content,
            'status': document.status,
            'user_id': current_user["id"]
        })
        return created or document
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/documents/{document_id}")
def delete_document(document_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a document"""
    # DB will enforce ownership; return generic message
    hybrid_service.delete_document(document_id)
    return {"message": "Document deleted successfully"}

@app.get("/api/documents/search/{query}")
def search_documents(query: str, current_user: dict = Depends(get_current_user)):
    """Search documents by name or content"""
    documents = hybrid_service.get_all_documents(user_id=current_user["id"])
    query_lower = query.lower()
    
    matching_documents = []
    for document in documents:
        if (query_lower in document.name.lower() or 
            query_lower in (document.content or "").lower()):
            matching_documents.append(document)
    
    return matching_documents

@app.get("/api/search")
async def search_papers(q: str = Query(...), top_k: int = 5, source: str = "semantic_scholar"):
    """
    Search for papers using multiple academic APIs.
    Sources: semantic_scholar, openalex, arxiv, pubmed
    """
    results = []

    # Helper to normalize date strings for sorting
    def to_sortable_date(date_str: str):
        try:
            if not date_str:
                return None
            # Try full ISO date first
            if len(date_str) >= 10 and date_str[4] == '-' and date_str[7] == '-':
                return datetime.fromisoformat(date_str[:10])
            # Try year-only
            if len(date_str) == 4 and date_str.isdigit():
                return datetime(int(date_str), 1, 1)
            # Fallback: attempt generic parse of leading date
            return datetime.fromisoformat(date_str.split('T')[0])
        except Exception:
            return None
    
    if source == "semantic_scholar" or source == "all":
        # Semantic Scholar API
        try:
            url = f"https://api.semanticscholar.org/graph/v1/paper/search?query={q}&limit={top_k}&fields=title,url,abstract,year,publicationDate"
            async with httpx.AsyncClient() as client:
                r = await client.get(url)
                if r.status_code == 200:
                    data = r.json()
                    for paper in data.get("data", []):
                        abstract = paper.get("abstract") or ""
                        title = paper.get("title") or "Untitled"
                        link = paper.get("url") or ""
                        year = paper.get("year")
                        pub_date = paper.get("publicationDate")  # ISO date string if provided
                        normalized_date = pub_date or (str(year) if year else "")
                        results.append({
                            "title": title,
                            "summary": abstract[:500],
                            "link": link,
                            "source": "Semantic Scholar",
                            "date": normalized_date
                        })
        except Exception as e:
            print(f"Semantic Scholar error: {e}")
    
    if source == "openalex" or source == "all":
        # OpenAlex API - following their best practices
        try:
            import urllib.parse
            encoded_query = urllib.parse.quote(q)
            # Use proper email for polite pool and better rate limits
            url = f"https://api.openalex.org/works?search={encoded_query}&per_page={top_k}&mailto=readnest@example.com"
            print(f"OpenAlex URL: {url}")
            async with httpx.AsyncClient() as client:
                r = await client.get(url)
                print(f"OpenAlex response status: {r.status_code}")
                if r.status_code == 200:
                    data = r.json()
                    print(f"OpenAlex found {len(data.get('results', []))} results")
                    for work in data.get("results", []):
                        # Extract abstract - OpenAlex uses abstract_inverted_index format
                        abstract_text = ""
                        if work.get("abstract_inverted_index"):
                            # Convert inverted index to readable text
                            abstract = work["abstract_inverted_index"]
                            words = []
                            for word, positions in abstract.items():
                                for pos in positions:
                                    words.append((pos, word))
                            words.sort()
                            abstract_text = " ".join([word for pos, word in words])
                        elif work.get("abstract"):
                            # Fallback to direct abstract if available
                            abstract_text = work["abstract"]
                        
                        title = work.get("title", "Untitled")
                        
                        # Get best available link following OpenAlex priority
                        link = ""
                        if work.get("open_access", {}).get("oa_url"):
                            # Open access PDF link (best option)
                            link = work["open_access"]["oa_url"]
                        elif work.get("primary_location", {}).get("landing_page_url"):
                            # Publisher landing page
                            link = work["primary_location"]["landing_page_url"]
                        elif work.get("doi"):
                            # DOI link
                            link = f"https://doi.org/{work['doi']}"
                        elif work.get("id"):
                            # OpenAlex page as fallback
                            link = work["id"]
                        
                        # Get publication year and venue information
                        pub_year = work.get("publication_year", "")
                        venue = work.get("primary_location", {}).get("source", {}).get("display_name", "")
                        
                        # Create enhanced title with year and venue
                        title_parts = [title]
                        if pub_year:
                            title_parts.append(f"({pub_year})")
                        if venue:
                            title_parts.append(f"- {venue}")
                        
                        enhanced_title = " ".join(title_parts)
                        
                        results.append({
                            "title": enhanced_title,
                            "summary": abstract_text[:500] if abstract_text else "No abstract available",
                            "link": link,
                            "source": "OpenAlex",
                            "date": str(pub_year) if pub_year else ""
                        })
                else:
                    print(f"OpenAlex error: {r.status_code} - {r.text}")
        except Exception as e:
            print(f"OpenAlex error: {e}")
            import traceback
            traceback.print_exc()
    
    if source == "arxiv" or source == "all":
        # arXiv API (no key required) - using correct URL format
        try:
            # Use HTTPS and proper URL encoding for arXiv API
            import urllib.parse
            encoded_query = urllib.parse.quote(f"all:{q}")
            url = f"https://export.arxiv.org/api/query?search_query={encoded_query}&start=0&max_results={top_k}&sortBy=relevance&sortOrder=descending"
            print(f"arXiv URL: {url}")
            async with httpx.AsyncClient() as client:
                r = await client.get(url)
                print(f"arXiv response status: {r.status_code}")
                if r.status_code == 200:
                    # Parse XML response using the correct namespace
                    from xml.etree import ElementTree as ET
                    root = ET.fromstring(r.text)
                    
                    # Find all entries using the correct namespace
                    entries = root.findall('.//{http://www.w3.org/2005/Atom}entry')
                    print(f"arXiv found {len(entries)} results")
                    
                    for entry in entries:
                        # Extract title
                        title_elem = entry.find('.//{http://www.w3.org/2005/Atom}title')
                        title = title_elem.text.strip() if title_elem is not None and title_elem.text else "Untitled"
                        
                        # Extract summary/abstract
                        summary_elem = entry.find('.//{http://www.w3.org/2005/Atom}summary')
                        summary = summary_elem.text.strip() if summary_elem is not None and summary_elem.text else "No abstract available"
                        
                        # Extract arXiv ID and create proper link
                        id_elem = entry.find('.//{http://www.w3.org/2005/Atom}id')
                        arxiv_id = ""
                        if id_elem is not None and id_elem.text:
                            # Extract arXiv ID from URL like "http://arxiv.org/abs/hep-ex/0307015"
                            arxiv_id = id_elem.text.split('/')[-1]
                        
                        # Create proper arXiv link
                        link = f"http://arxiv.org/abs/{arxiv_id}" if arxiv_id else ""
                        
                        # Get publication date
                        published_elem = entry.find('.//{http://www.w3.org/2005/Atom}published')
                        pub_date = ""
                        if published_elem is not None and published_elem.text:
                            # Prefer full date (YYYY-MM-DD) if present
                            full_date = published_elem.text[:10]
                            pub_date = full_date
                        year_suffix = f" ({pub_date[:4]})" if pub_date else ""
                        
                        # Get authors
                        authors = []
                        for author in entry.findall('.//{http://www.w3.org/2005/Atom}author'):
                            name_elem = author.find('.//{http://www.w3.org/2005/Atom}name')
                            if name_elem is not None and name_elem.text:
                                authors.append(name_elem.text.strip())
                        
                        author_text = f" by {', '.join(authors[:3])}" if authors else ""
                        if len(authors) > 3:
                            author_text += f" et al."
                        
                        results.append({
                            "title": title + year_suffix + author_text,
                            "summary": summary[:500],
                            "link": link,
                            "source": "arXiv",
                            "date": pub_date
                        })
                else:
                    print(f"arXiv error: {r.status_code} - {r.text}")
        except Exception as e:
            print(f"arXiv error: {e}")
            import traceback
            traceback.print_exc()
    
    if source == "pubmed" or source == "all":
        # PubMed API (no key required for basic search)
        try:
            # First search for PMIDs
            import urllib.parse
            encoded_query = urllib.parse.quote(q)
            search_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi?db=pubmed&term={encoded_query}&retmax={top_k}&retmode=json"
            print(f"PubMed search URL: {search_url}")
            async with httpx.AsyncClient() as client:
                r = await client.get(search_url)
                if r.status_code == 200:
                    search_data = r.json()
                    pmids = search_data.get("esearchresult", {}).get("idlist", [])
                    print(f"PubMed found {len(pmids)} PMIDs")
                    
                    if pmids:
                        # Get details for each PMID
                        pmids_str = ",".join(pmids)
                        fetch_url = f"https://eutils.ncbi.nlm.nih.gov/entrez/eutils/efetch.fcgi?db=pubmed&id={pmids_str}&retmode=xml"
                        
                        r2 = await client.get(fetch_url)
                        if r2.status_code == 200:
                            # Parse XML response
                            from xml.etree import ElementTree as ET
                            root = ET.fromstring(r2.text)
                            
                            articles = root.findall('.//PubmedArticle')
                            for article in articles:
                                title_elem = article.find('.//ArticleTitle')
                                title = title_elem.text if title_elem is not None else "Untitled"
                                
                                abstract_elem = article.find('.//AbstractText')
                                abstract = abstract_elem.text if abstract_elem is not None else "No abstract available"
                                
                                pmid_elem = article.find('.//PMID')
                                pmid = pmid_elem.text if pmid_elem is not None else ""
                                link = f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/" if pmid else ""
                                
                                # Get publication date (best-effort)
                                year_elem = article.find('.//PubDate/Year')
                                month_elem = article.find('.//PubDate/Month')
                                day_elem = article.find('.//PubDate/Day')
                                pub_year = year_elem.text if year_elem is not None else ""
                                # Normalize month text (could be Jan/01/etc.)
                                month = month_elem.text if month_elem is not None else "01"
                                day = day_elem.text if day_elem is not None else "01"
                                # Map short month names to numbers
                                month_map = {"Jan":"01","Feb":"02","Mar":"03","Apr":"04","May":"05","Jun":"06","Jul":"07","Aug":"08","Sep":"09","Oct":"10","Nov":"11","Dec":"12"}
                                month_norm = month_map.get(month, month.zfill(2)) if pub_year else ""
                                pub_date = f"{pub_year}-{month_norm}-{day.zfill(2)}" if pub_year else ""
                                year_suffix = f" ({pub_year})" if pub_year else ""
                                
                                results.append({
                                    "title": title + year_suffix,
                                    "summary": abstract[:500],
                                    "link": link,
                                    "source": "PubMed",
                                    "date": pub_date or (pub_year or "")
                                })
        except Exception as e:
            print(f"PubMed error: {e}")
            import traceback
            traceback.print_exc()

    # Sort results by date descending (newest first); items without date go last
    results.sort(key=lambda r: (to_sortable_date(r.get("date")) is None, to_sortable_date(r.get("date")) or datetime.min), reverse=False)
    results.reverse()

    return {"results": results}

@app.post("/api/scholar-agent")
async def run_scholar_agent(payload: dict):
    """
    payload: {"user_prompt": "...", "papers": [...]}
    """
    try:
        # invoke the compiled graph
        result = scholar_agent.invoke({
            "user_prompt": payload.get("user_prompt", ""),
            "papers": payload.get("papers", [])
        })

        # result should be a dict with 'results'
        if isinstance(result, dict):
            results = result.get("results", [])
            # return normalized response (include error/raw if present)
            return {"results": results, "error": result.get("error"), "raw": result.get("raw")}
        else:
            # try parse if it's a string (rare)
            try:
                parsed = json.loads(result)
                return {"results": parsed if isinstance(parsed, list) else [], "raw": result}
            except Exception:
                return {"results": [], "raw": str(result)}

    except Exception as e:
        tb = traceback.format_exc()
        print("scholar_agent endpoint error:", tb)
        raise HTTPException(status_code=500, detail=str(e))

class ChatRequest(BaseModel):
    message: str
    conversation_history: List[dict] = []

@app.post("/api/chat")
async def chat_with_ai(request: ChatRequest, current_user: dict = Depends(get_current_user)):
    """
    AI chat endpoint with access to user's notes/journals for context
    """
    try:
        from langchain_groq import ChatGroq
        
        # Initialize Groq client using langchain_groq (same as scholar_agent)
        llm = ChatGroq(
            model="llama-3.1-8b-instant",
            temperature=0.7,
            groq_api_key=os.getenv("GROQ_API_KEY"),
        )
        
        # Fetch user's journal entries for context
        notes_context = ""
        try:
            journals = hybrid_service.get_all_journals(user_id=current_user["id"])
            if journals:
                # Include recent notes (last 5 entries, truncated)
                recent_notes = journals[:5]
                notes_summary = "\n".join([
                    f"- {entry.title}: {entry.content[:200]}..." if len(entry.content) > 200 else f"- {entry.title}: {entry.content}"
                    for entry in recent_notes
                ])
                notes_context = f"\n\nYou have access to the user's notes/journal entries. Here are their recent notes:\n{notes_summary}\n"
        except Exception as e:
            print(f"Error fetching notes for context: {e}")
            notes_context = ""
        
        # Build conversation context
        conversation_context = "You are a helpful AI assistant for ReadNest, a reading and research platform. You can help with general questions, explanations, creative tasks, and casual conversation. " + \
                              "You also have access to the user's notes and journal entries, which you can reference when relevant to answer their questions." + \
                              "Be friendly, informative, and concise.\n\n"
        
        # Add notes context if available
        if notes_context:
            conversation_context += notes_context
        
        # Add conversation history (last 10 messages)
        for msg in request.conversation_history[-10:]:
            if msg.get('role') == 'user':
                conversation_context += f"User: {msg['content']}\n"
            elif msg.get('role') == 'assistant':
                conversation_context += f"Assistant: {msg['content']}\n"
        
        # Add current message
        conversation_context += f"User: {request.message}\nAssistant:"
        
        # Get response from Groq
        response = llm.invoke(conversation_context)
        ai_response = response.content
        
        return {"response": ai_response}
        
    except Exception as e:
        tb = traceback.format_exc()
        print("chat endpoint error:", tb)
        raise HTTPException(status_code=500, detail=str(e))

# Auth endpoints (Supabase)
@app.post("/auth/register")
def register_user(payload: RegisterRequest):
    try:
        res = supabase.auth.sign_up({
            "email": payload.email,
            "password": payload.password
        })
        return {
            "user": {"id": res.user.id, "email": res.user.email} if res.user else None,
            "session": {
                "access_token": getattr(res.session, "access_token", None),
                "refresh_token": getattr(res.session, "refresh_token", None)
            } if res.session else None
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/auth/login")
def login_user(payload: LoginRequest):
    try:
        res = supabase.auth.sign_in_with_password({
            "email": payload.email,
            "password": payload.password
        })
        if not res.session:
            raise HTTPException(status_code=400, detail="Invalid credentials")
        return {
            "access_token": res.session.access_token,
            "refresh_token": res.session.refresh_token,
            "user": {"id": res.user.id, "email": res.user.email} if res.user else None
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))

@app.post("/auth/logout")
def logout_user():
    # Stateless API: client should discard its token. This endpoint exists for symmetry.
    try:
        # Attempt to sign out (no-op without a persisted session in this server)
        try:
            supabase.auth.sign_out()
        except Exception:
            pass
        return {"success": True}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))
